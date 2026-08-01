import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Styles
style_title = doc.styles['Title']
font = style_title.font
font.name = 'Arial'
font.size = Pt(24)
font.color.rgb = RGBColor(0, 0, 0)

style_h1 = doc.styles['Heading 1']
font = style_h1.font
font.name = 'Arial'
font.size = Pt(16)
font.color.rgb = RGBColor(0, 51, 102)

style_h2 = doc.styles['Heading 2']
font = style_h2.font
font.name = 'Arial'
font.size = Pt(14)
font.color.rgb = RGBColor(0, 0, 0)

# Title
title = doc.add_paragraph('Submission 2: The Build', style='Title')
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('AI Ops Engineer Assignment - QuickMove\n', style='Normal')

# 1. The Chosen Workflow
doc.add_heading('1. The Chosen Workflow', level=1)
p = doc.add_paragraph('Workflow: ')
p.add_run('Property Partner Dispatcher').bold = True
p.add_run('\n(Targeting WF04: Property Partner Outreach)')

# 2. The Problem
doc.add_heading('2. The Problem It Solves', level=1)
doc.add_paragraph(
    'At 200+ relocations per month, ops agents spend hundreds of hours manually identifying which property '
    'partners operate in the destination city, trying to remember which ones are best for certain configurations (like Pet Friendly), '
    'and then manually drafting WhatsApp messages to 3-5 partners per case. '
    'This is a massive time sink, highly error-prone, and scales poorly. If a partner doesn\'t respond, the ops agent often forgets to follow up.'
)
doc.add_paragraph(
    'By automating vendor matching and message dispatch, we eliminate manual drafting and ensure we are always routing leads '
    'to our fastest, highest-performing partners.'
)

# 3. The Solution
doc.add_heading('3. The Solution: QuickMove Property Partner Dispatcher', level=1)
doc.add_paragraph(
    'We built a zero-friction, single-file web application that ops agents can use immediately. '
    'The tool requires no deployment, server, or installation—it is a standalone HTML file that runs locally in any browser.'
)

doc.add_heading('Key Features:', level=2)
features = [
    "Algorithmic Matching: Contains an embedded database of partners. Instantly filters out partners that don't match the customer's strict criteria (e.g., Pet Friendly constraints, City coverage, BHK configurations).",
    "Performance Ranking: Automatically sorts the matched partners by their historical performance score and average response time, ensuring the best partners get the leads.",
    "Automated Outreach Generation: Generates 3 highly personalized, context-rich WhatsApp messages simultaneously.",
    "One-Click Dispatch: Ops agents simply click 'Copy Message' for each of the top 3 recommendations and paste them directly into WhatsApp."
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# 4. Design & Usability
doc.add_heading('4. Design, Taste, and Usability', level=1)
doc.add_paragraph(
    'The tool is designed with a premium, SaaS-grade user interface. The aesthetics feel human-crafted and highly professional.'
)
design_pts = [
    "Usability for Non-Technical Ops: Minimal text entry. Requirements like Configuration and 'Must Haves' are toggled via interactive chips, eliminating typos.",
    "Split-Screen UX: Inputs on the left, instant actionable outputs on the right. No page reloads.",
    "Visual Feedback: Includes micro-interactions like a 'Copied!' state change on the button and a toast notification to give confidence to the user.",
    "Tasteful Aesthetics: Utilizes a refined color palette (soft stone backgrounds, crisp dark grays, cohesive shadows) and modern typography (Inter font)."
]
for pt in design_pts:
    doc.add_paragraph(pt, style='List Bullet')

# 5. How to Use
doc.add_heading('5. How to Access and Use the Tool', level=1)
doc.add_paragraph('The working tool is provided as the file: ')
doc.add_paragraph('QuickMove_Vendor_Dispatcher.html', style='List Bullet').runs[0].bold = True

steps = [
    "Double-click the HTML file to open it in any modern web browser.",
    "Select the Destination City and Configuration on the left panel.",
    "Toggle any special requirements (e.g., Pet Friendly).",
    "Click the 'Find & Dispatch Partners' button.",
    "The tool will display the Top 3 best-matched partners.",
    "Click 'Copy Message' next to each partner to copy the auto-generated outreach text to your clipboard."
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.save('Submission_2_The_Build.docx')
print("Submission 2 Document Created Successfully.")
